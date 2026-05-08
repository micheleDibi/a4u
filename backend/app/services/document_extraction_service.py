"""Estrazione testo da documenti caricati (PDF/DOCX/DOC/RTF/TXT/MD).

Tutte le librerie usate (`pdfplumber`, `python-docx`, `docx2txt`, `striprtf`)
sono pure-Python o wheel-installabili su Windows. La extraction è blocking
quindi viene wrappata in `asyncio.to_thread`.

Solleva `DocumentExtractionError(message)` quando il file è corrotto,
protetto da password, o di formato non supportato.
"""
from __future__ import annotations

import asyncio
from pathlib import Path

from app.core.config import get_settings
from app.core.logging import get_logger

log = get_logger("app.document_extraction")


class DocumentExtractionError(Exception):
    """Errore durante l'estrazione del testo da un documento."""


# Lista MIME → extractor (sync). Le estensioni file sono fallback
# se il MIME risulta vuoto o non riconosciuto.
PDF_MIMES = {"application/pdf"}
DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
DOC_MIMES = {"application/msword"}
RTF_MIMES = {"application/rtf", "text/rtf"}
TEXT_MIMES = {"text/plain", "text/markdown"}


def _extract_pdf(path: Path) -> str:
    import pdfplumber

    parts: list[str] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text:
                    parts.append(text)
    except Exception as exc:
        msg = str(exc).lower()
        if "password" in msg or "encrypt" in msg:
            raise DocumentExtractionError(
                "Il PDF è protetto da password e non può essere letto."
            ) from exc
        raise DocumentExtractionError(
            f"Impossibile leggere il PDF: {exc}"
        ) from exc
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    import docx

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise DocumentExtractionError(
            f"Impossibile leggere il file DOCX: {exc}"
        ) from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_doc(path: Path) -> str:
    import docx2txt

    try:
        text = docx2txt.process(str(path))
    except Exception as exc:
        raise DocumentExtractionError(
            f"Impossibile leggere il file DOC: {exc}"
        ) from exc
    return text or ""


def _extract_rtf(path: Path) -> str:
    from striprtf.striprtf import rtf_to_text

    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
        return rtf_to_text(raw)
    except Exception as exc:
        raise DocumentExtractionError(
            f"Impossibile leggere il file RTF: {exc}"
        ) from exc


def _extract_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        raise DocumentExtractionError(
            f"Impossibile leggere il file di testo: {exc}"
        ) from exc


def _dispatch(path: Path, mime_type: str) -> str:
    mime = (mime_type or "").lower()
    suffix = path.suffix.lower()
    if mime in PDF_MIMES or suffix == ".pdf":
        return _extract_pdf(path)
    if mime in DOCX_MIMES or suffix == ".docx":
        return _extract_docx(path)
    if mime in DOC_MIMES or suffix == ".doc":
        return _extract_doc(path)
    if mime in RTF_MIMES or suffix == ".rtf":
        return _extract_rtf(path)
    if mime in TEXT_MIMES or suffix in {".txt", ".md", ".markdown"}:
        return _extract_text_file(path)
    raise DocumentExtractionError(
        f"Formato non supportato: mime='{mime_type}', estensione='{suffix}'."
    )


async def extract_text(file_path: Path, mime_type: str) -> tuple[str, int]:
    """Estrae il testo da `file_path`. Ritorna `(text_truncated, original_chars)`.

    Il testo viene troncato a `settings.course_document_max_chars` per
    contenere il consumo di token. `original_chars` contiene la lunghezza
    pre-troncamento (utile per logging).
    """
    if not file_path.exists():
        raise DocumentExtractionError(f"File non trovato: {file_path}")

    settings = get_settings()
    raw = await asyncio.to_thread(_dispatch, file_path, mime_type)
    raw = (raw or "").strip()
    original = len(raw)
    if not raw:
        raise DocumentExtractionError(
            "Documento privo di testo estraibile (forse è una scansione? "
            "OCR non supportato in questa versione)."
        )

    max_chars = max(1000, int(settings.course_document_max_chars))
    if original > max_chars:
        log.warning(
            "course_document_text_truncated",
            path=str(file_path),
            original=original,
            kept=max_chars,
        )
        raw = raw[:max_chars]
    return raw, original
